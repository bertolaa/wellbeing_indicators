import streamlit as st
import pandas as pd
import math
import json
import requests
from pathlib import Path
import altair as alt
import xlrd
import openai as client
import toml
from docx import Document
from docx.shared import Inches
from io import BytesIO
import kaleido

#-----------------------------------------------------------------
# Step 1: Get OpenAI API key
#-----------------------------------------------------------------
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
client.api_key = OPENAI_API_KEY


# Set the title and favicon that appear in the Browser's tab bar.
st.set_page_config(
    page_title='Health in well-being economy analysis tool',
    layout="wide",
)

# Display image as heading
st.image(Path(__file__).parent/'img/logo.gif')


# Declare some useful functions.
doc = Document()

@st.cache_data
def get_wb_data(url):
    # Fetch data from the World Bank API in JSON format
    response = requests.get(url)
    datalist = response.json()    
    
    print ("url WB-->", url_a)
    # Extract relevant data
    years = []
    gdp_values = []
    country_codes = []
    
    for entry in datalist[1]:
        years.append(entry['date'])
        gdp_values.append(entry['value'])
        country_codes.append(entry['countryiso3code'])

    # Convert to DataFrame
    gdp_df = pd.DataFrame({
        'Country Code': country_codes,
        'Year': years,
        'Value': gdp_values
    })
        
    MIN_YEAR = 1960
    MAX_YEAR = 2024

    # Convert years from string to integers
    gdp_df['Year'] = pd.to_numeric(gdp_df['Year'])
    
    print (gdp_df.head())
    return gdp_df

def get_data_from_eurostat (url_code):
    import eurostat
    global filter_list, sex_split

    #getting dimensions and data as pandas df
    df = eurostat.get_data_df(url_code)
    cols = list(df.columns)

    #setting up filter string for unpivoting
    not_unique_col = []
    unpivot_string = []
    vars_string = []
    time_period = False

    #defining vars to unpivot
    for i in cols:
            if (df[i].nunique() >1) or (): 
                not_unique_col.append(i)

    for i in not_unique_col:
        if (time_period):
            unpivot_string.append(i)
        else:
            vars_string.append(i)
            if i[:3] == 'geo':
                time_period = True
                toberenamed = i 
            else:
                unique_values = df[i].unique()
                
                if i == 'sex':
                    sex_split = True
                else: 
                    filter_list[i] = unique_values.tolist()
    
    #print ("id_vars", vars_string)
    #print ("value_vars", unpivot_string)                
    gdp_df = pd.melt (df, id_vars = vars_string, value_vars= unpivot_string, var_name='Year', value_name="Value")
    gdp_df = gdp_df.rename (columns = {toberenamed : 'Country Code'})
    gdp_df['Year'] = pd.to_numeric(gdp_df['Year'])
    return gdp_df

def get_data_from_whoeurope (url_code):
    global filter_list, sex_split

    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
    response = requests.get(url_code, headers=headers)

    dt = response.json()
    #print ("JSON Response status reading WHO api", response.status_code)

    sex_split = False
    df_string = {}
    years = []
    values = []
    country_codes = []
    sex = []
    dim = []

    for i in (dt[0]['dimensions']):
        dim.append(i['code'])

    if ("SEX" in dim): sex_split = True
        
    for entry in dt[0]['data']:
        years.append(entry['dimensions']['YEAR'])
        values.append(entry['value']['numeric'])
        country_codes.append(entry['dimensions']['COUNTRY'])   
        df_string = {'Country Code': country_codes, 'Year': years, 'Value': values}

        if (sex_split): 
            sex.append(entry['dimensions']['SEX'])
            df_string.update ({'sex': sex})

    # Convert to DataFrame
    df = pd.DataFrame(df_string)

    #check whether dataset has disaggregation by sex
    if (sex_split):
        if (df['sex'].nunique() < 2): sex_split = False

    df_cleaned = df[ df['Country Code'] != '']
    gdp_df = df_cleaned.sort_values(by=['Country Code', 'Year'])
    gdp_df['Year'] = pd.to_numeric(gdp_df['Year'])
    print (gdp_df.columns)
    return gdp_df

def get_data_from_OECD (url_code):
    global filter_list, sex_split

    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
    response = requests.get(url_code, headers=headers)
    
    print ("JSON Response status reading WHO api", response.status_code)

    dt = response.json()   
    
    df_string = {}
    years = []
    values = []
    country_codes = []
    sex = []
    dim = []

    for i in (dt[0]['SeriesKey']):
        dim.append(i['REF_AREA'])
    
    
    for entry in dt[0]['data']:
        years.append(entry['dimensions']['YEAR'])
        values.append(entry['value']['numeric'])
        country_codes.append(entry['dimensions']['COUNTRY'])   
        df_string = {'Country Code': country_codes, 'Year': years, 'Value': values}

        if (sex_split): 
            sex.append(entry['dimensions']['SEX'])
            df_string.update ({'sex': sex})

    # Convert to DataFrame
    df = pd.DataFrame(df_string)

    if (sex_split):
        if (df['sex'].nunique() < 2): sex_split = False

    df_cleaned = df[ df['Country Code'] != '']
    gdp_df = df_cleaned.sort_values(by=['Country Code', 'Year'])
    gdp_df['Year'] = pd.to_numeric(gdp_df['Year'])
    return gdp_df

def get_data_from_WHOHESR (url_code):
    global sex_split, filter_list
    
    not_unique_col = []
    unpivot_string = []
    vars_string = []
    
    if (source== "WHO/HESRI"): data_filename = Path(__file__).parent/'data/HESR1.xlsx'
    else: data_filename = Path(__file__).parent/'data/HESR2.xlsx'
    
    hesr = pd.read_excel(data_filename)                     
    df = pd.DataFrame(hesr)  
    print ("len(df) before filtering",len(df))
                    
    df = df[df['indicator_abbr'] == url_code]
    print ("len(df) after filtering",len(df))

    cols = list(df.columns)
    
    for i in cols:
        if (df[i].nunique()>1): not_unique_col.append(i)
       
    for i in not_unique_col:
        vars_string.append(i)
        unique_values = df[i].unique()
                
        if i == 'sex':
            sex_split = True
        else: 
            filter_list[i] = unique_values.tolist()
    
    
    #lighten the dataframe        
    keys_to_remove = ['Country Code', 'Year', 'Value', 'population']  
    if (source == 'WHO/HESRI 2'): keys_to_remove += ['Education', 'Income']
    for key in keys_to_remove: 
        if key in filter_list: del filter_list[key]
    
    #df = df.loc[df.groupby('Country Code')['Year'].idxmax()]
    
    df_cleaned = df.dropna()  
    print("filter_list", filter_list)
    print (df.head())    
    return df_cleaned

def draw_chart (df, measure, container, sexdim):
    global sex_split
    filter = 'Country Code:N'

    chart = alt.Chart(df).mark_line(point=True).encode(
                x=alt.X('Year:Q', sort='ascending', axis=alt.Axis(format='d')),   
                y='Value:Q',
                color= filter,
                tooltip=['Year', 'Value', 'Country Code']
            ).properties().interactive()
            
    if sex_split: headertext = measure + " - Sex = " + sexdim
    else:  headertext = measure

    pivot_data = df.pivot_table(index="Country Code", columns="Year", values="Value").round(2)
    
    with container:
        st.header (headertext)
        st.altair_chart(chart, use_container_width=True)        
        st.header ("Data table - " + headertext)   
        st.write(pivot_data)
    return

def draw_chart_hesr (df, measure, container, sexdim, dim):
    global sex_split, filter_criteria
    filter = 'Country Code:N'
    dim_c = dim + ":N"

    # Create the line chart with separate lines for each country
    lines = alt.Chart(df).mark_line(color='black').encode(
        alt.X('Value:Q'),
        alt.Y('Country Code:N', sort=alt.EncodingSortField(field='Country Code', order='ascending')),
        alt.Detail('Country Code:N')  # Create separate lines for each country
    )
    
    # Create the point chart
    points = alt.Chart(df).mark_point(filled=True, size=120).encode(
        alt.X('Value:Q'),
        alt.Y('Country Code:N', sort=alt.EncodingSortField(field='Country', order='ascending')),
        #alt.Color( dim_c, legend=alt.Legend(title=dim_c, values=['Low', 'Medium', 'High']))  # Specify legend order
        alt.Color( 'Attributes:N', legend=alt.Legend(title=dim, values=['Low', 'Medium', 'High']))  # Specify legend order
    )
    
    # Combine the charts
    chart = (points + lines).properties(title=dim + ' Levels by Country')
            
    if sex_split: headertext = measure + " - Sex = " + sexdim
    else:  headertext = measure
    
    
    #to create list of stratifiers in the table
    filt = "Country Code," + dim    
    idx = filt.split(',') 
   
    print ("idx:", idx)
    pivot_data = df.pivot_table( index= ['Country Code','Attributes'],  columns="Year", values="Value").round(2)
    
    with container:
        st.header (headertext)
        st.altair_chart(chart, use_container_width=True)        
        st.header ("Data table - " + headertext)   
        st.write(pivot_data)
    return

##############################################################################################################
# Draw the actual page
# Set the title that appears at the top of the page.
##############################################################################################################

st.header(":green[Health in well-being economy analysis tool]")
#st.write("_... to visualize well-being data from public databases_")

# Populating variables on "indicators.xls" and "countries_WHO_Euro.csv" files
data_filename = Path(__file__).parent/'data/Indicators.xlsx'
indicators = pd.read_excel(data_filename)                 
indicators_df = pd.DataFrame(indicators)                
sex_split = False
filter_list = {}

countries_WHOEURO = pd.read_excel(Path(__file__).parent/'data/countries_WHO_Euro.xls') 
countries_df = pd.DataFrame(countries_WHOEURO)

#Picking ISO3 as country code to match data
countries_df['Country Code'] = countries_df['Countries.code']

mod = "Country Profile" 
options = ["Explore Countries", "Country Profile"]    

#mod = st.pills ("Select type of visualization", options) 

#Explore countries 
if ( mod == "Explore countries" ):
    # Two equal columns:
    col1, col2 = st.columns(2)

    #container = st.container(border=True)
    c1 = col1.container(border=True)
    c1.subheader ("Main indicators / countries selection")

    # Initialize session state for category/indicator/country inputs
    if 'selected_datasource' not in st.session_state:
        st.session_state.selected_datasource = indicators_df['Indicator.datasource'].unique()[0]
    if 'selected_cat' not in st.session_state:
        st.session_state.selected_cat = indicators_df[indicators_df['Indicator.datasource'] == st.session_state.selected_datasource]['Indicator.area'].unique()[0]
    if 'selected_ind' not in st.session_state:
        st.session_state.selected_ind = indicators_df[(indicators_df['Indicator.datasource'] == st.session_state.selected_datasource) & 
                                            (indicators_df['Indicator.area'] == st.session_state.selected_cat)]['Indicator.short_name'].unique()[0]

    # Function to update the options for Type based on selected Category
    def update_cat():
        st.session_state.selected_cat = indicators_df[indicators_df['Indicator.datasource'] == st.session_state.selected_datasource]['Indicator.area'].unique()[0]
        update_ind()

    # Function to update the options for Item based on selected Type
    def update_ind():
        st.session_state.selected_ind= indicators_df[(indicators_df['Indicator.datasource'] == st.session_state.selected_datasource) & 
                                            (indicators_df['Indicator.area'] == st.session_state.selected_cat)]['Indicator.short_name'].unique()[0]
       
    # Create widgets for each layer of input
    c1.pills(
        ''':green[**1/5 - Data from which datasource?**]''', 
        options = indicators_df['Indicator.datasource'].unique(), 
        key='selected_datasource',
        on_change=update_cat
    )

    c1.selectbox(
        ''':green[**2/5 Select a category**]''',
        options=indicators_df[indicators_df['Indicator.datasource'] == st.session_state.selected_datasource]['Indicator.area'].unique(),
        key='selected_cat',
        on_change=update_ind
    )

    c1.selectbox(
        ''':green[**3/5 Select an indicator**]''',
        options=indicators_df[(indicators_df['Indicator.datasource'] == st.session_state.selected_datasource) & 
                (indicators_df['Indicator.area'] == st.session_state.selected_cat)]['Indicator.short_name'].unique(),
        key='selected_ind'
    )

    source = st.session_state.selected_datasource
    url_code = ''.join(indicators_df[indicators_df['Indicator.short_name'] == st.session_state.selected_ind]['Indicator_Code'].values[0])
    measure = indicators_df[indicators_df['Indicator.short_name'] == st.session_state.selected_ind]['Indicator.short_name'].values[0]
    ind_longtitle = indicators_df[indicators_df['Indicator.short_name'] == st.session_state.selected_ind]['Indicator.long_name'].values[0]

    #selecting data source
    match source:
        case "OECD":
            url_a = "https://sdmx.oecd.org/public/rest/data/" + url_code
            url_b = "https://data-explorer.oecd.org/vis?df[ds]=dsDisseminateFinalDMZ&df[id]=" + url_code +"&df[ag]=OECD.ELS.HD"  
            gdp_df = get_data_from_OECD(url_a)
                    
        case "WORLD BANK":
            url_a = "https://api.worldbank.org/v2/country/all/indicator/"+ url_code +"?format=json&per_page=20000"
            url_b = "https://data.worldbank.org/indicator/" + url_code
            gdp_df = get_wb_data(url_a)  
            
        case "EUROSTAT":
            url_b = "https://ec.europa.eu/eurostat/web/products-datasets/-/"+ url_code
            gdp_df = get_data_from_eurostat(url_code)
            countries_df['Country Code'] = countries_df['Countries.iso2']
        
        case "WHO/Europe":
            url_a = "https://dw.euro.who.int/api/v3/Batch/Measures?codes="+ url_code
            url_b = "https://dw.euro.who.int/api/v3/Batch/Measures?codes="+ url_code
            gdp_df = get_data_from_whoeurope(url_a)

        case "WHO/HESRI":
            url_a = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
            url_b = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
            gdp_df = get_data_from_WHOHESR (url_code)
            
        case "WHO/HESRI 2":
            url_a = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
            url_b = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
            gdp_df = get_data_from_WHOHESR (url_code)
        
        case _:
            c1.write ("No datasource selected")

    countries_df = countries_df.set_index('Country Code')

    c = gdp_df['Country Code'].unique()   
    cc = pd.DataFrame(c, columns=['Country Code'])              

    #extract only Countries data from WHO/EURO
    countries = pd.merge (cc, countries_df, on="Country Code" )                    

    gdp_df = gdp_df.dropna()  
    
    min_value = gdp_df['Year'].min()
    max_value = gdp_df['Year'].max()

    from_year, to_year = c1.slider(
        ''':green[**4/5 - Which years are you interested in?**]''',
        min_value=min_value,
        max_value=max_value,
        value=[min_value, max_value])

    if not len(countries):
        st.warning("Select at least one country")

    country_container = c1.container()  
    sci_selected = c1.checkbox("Small Countries Initiative")  

    #Countried selection pane
    if sci_selected:
        selected_countries = country_container.multiselect (''':green[**5/5 - Which countries would you like to view?**]''',
        countries['Countries.short_name'], countries.loc[countries['group']=='SCI', 'Countries.short_name'] )         
    else:
        selected_countries = country_container.multiselect (''':green[**5/5 - Which countries would you like to view?**]''',
        countries['Countries.short_name'])

    filtered_countries = countries[countries['Countries.short_name'].isin(selected_countries)]
    column_titles = filtered_countries.columns.tolist()
    iso_acronyms = filtered_countries['Country Code'].tolist()
    filter_criteria = {}

    c2  = col2.container(border=True)

    #panel for optional filters - if any - in the data source
    if len(filter_list) > 0:
        print("len filter list", len(filter_list))
        
        c2.subheader ("Filter dimensions...")    
        if sex_split: c2.write(":green[**- Sex disaggregation**]")       
        print (filter_list)
        
        #Filtering data with indicator-specific dimensions
        for key, values in filter_list.items():
            filter_criteria[key] = c2.selectbox(f':green[**Select a value for {key}**]', values)
            print("filter_criteria[key]", filter_criteria[key])
                
    elif sex_split: 
        c2.subheader ("Filter dimensions...")    
        c2.write(":green[**- Sex disaggregation**]")

    # Filter the data
    filtered_gdp_df = gdp_df[
            (gdp_df['Country Code'].isin(iso_acronyms)) & (gdp_df['Year'] <= to_year) & (from_year <= gdp_df['Year'])
        ]

    #aligning boxes
    col1, col2 = st.columns(2)
    c1 = col1.container(border=True)
    c2 = col2.container(border=True)

    if len(filtered_gdp_df)>0:
        data_to_chart = filtered_gdp_df
        
        #filtering based on previous selections
        if (source[:9] != 'WHO/HESRI'):
            for key, value in filter_criteria.items():            
                # filter with dimensions criteria selected in the c2 dropdown list 
                data_to_chart = data_to_chart[data_to_chart[key] == value]
        else:
                data_to_chart = data_to_chart[data_to_chart['dimension'] == filter_criteria['dimension']]          
                
#########################           
        if sex_split:
            #prep the chart for sex=F
            dtc = data_to_chart[(data_to_chart['sex'] == 'F') | (data_to_chart['sex'] == 'FEMALE')]
            if (source[:9]=='WHO/HESRI'): 
                # Group by country and get the index of the maximum year for each country
                dtc_ = dtc.groupby(['Country Code', 'Attributes'])['Year'].idxmax()
                dtc_latest = dtc.loc[dtc_]
                draw_chart_hesr(dtc_latest, measure, c1, 'F', filter_criteria['dimension'])
            else: draw_chart(dtc, measure, c1, 'F')
        
            #prep the chart for sex=M
            dtc = data_to_chart[(data_to_chart['sex'] == 'M')  | (data_to_chart['sex'] == 'MALE')]
            if (source=='WHO/HESRI 2'): 
                # Group by country and get the index of the maximum year for each country
                dtc_ = dtc.groupby(['Country Code', 'Attributes'])['Year'].idxmax()
                dtc_latest = dtc.loc[dtc_]
                draw_chart_hesr(dtc_latest, measure, c2, 'M', filter_criteria['dimension'])
            else: draw_chart(dtc, measure, c2, 'M')

        else:           
            #drawing the chart when sex is not a dimension
            draw_chart (data_to_chart, measure, c1, "")
                                
####################################                
    else:
        c1.write(":red[Select indicator and country... No data to display with this selection]")

    col1, col2 = st.columns(2)
    c1 = col1.container(border=False)
    c1.html("<a href="+url_b+" target='_blank'>Data source</a>")         
    
###############################################################################################
# selected the mode Country Profile from the main page
elif (mod == "Country Profile"):
    prompt0 = ' '
    col1, col2 = st.columns(2)
    c1 = col1.container(border=False)
   
    #country selection
    selected_country= c1.selectbox(
     ''':green[*Which WHO/Europe country would you like to select?*]''', countries_df['Countries.short_name'], index=None)

    filtered_countries = countries_df[countries_df['Countries.short_name'] == selected_country]       
    iso_acronyms = filtered_countries['Country Code'].to_list()
    indi_df = indicators_df[indicators_df['Country_Profile'] == True]
    
    #By clicking the button we start the production of the Country profile
    if ((c1.button ("Produce Country profile")) & (selected_country != None)):

        #start populating document to be downloaded
        doc.add_heading('Health in well-being economy analysis tool', 0)
        
        a = "Selected indicators to be elaborated: " + str(len(indi_df)) + " - Country: " + selected_country 
        c1.subheader (a)
        doc.add_paragraph(a)        
        
        #iterate all the indicators to be included in the report
        for index, row in indi_df.iterrows():
            iso_acronyms = filtered_countries['Country Code'].to_list()
            url_code = row['Indicator_Code']
            
            #title of the indicator     
            chart_title = str(row['Indicator.datasource'])+ " - " +(str(row['Indicator.short_name']))       
            c1.subheader (chart_title)
            doc.add_paragraph(chart_title, style="Heading 1")                  
                        
            source = row['Indicator.datasource']            
            #getting data
            match source:
                case "OECD":
                    url_a = "https://sdmx.oecd.org/public/rest/data/" + url_code
                    url_b = "https://data-explorer.oecd.org/vis?df[ds]=dsDisseminateFinalDMZ&df[id]=" + url_code +"&df[ag]=OECD.ELS.HD"  
                    gdp_df = get_data_from_OECD(url_a)
                            
                case "WORLD BANK":
                    url_a = "https://api.worldbank.org/v2/country/all/indicator/"+ url_code +"?format=json&per_page=20000"
                    url_b = "https://data.worldbank.org/indicator/" + url_code
                    gdp_df = get_wb_data(url_a)  
                    
                case "EUROSTAT":
                    url_b = "https://ec.europa.eu/eurostat/web/products-datasets/-/"+ url_code
                    gdp_df = get_data_from_eurostat(url_code)
                    iso_acronyms = filtered_countries['Countries.iso2'].to_list()
                    url_a = url_b
                                  
                case "WHO/Europe":
                    url_a = "https://dw.euro.who.int/api/v3/Batch/Measures?codes="+ url_code
                    url_b = "https://dw.euro.who.int/api/v3/Batch/Measures?codes="+ url_code
                    gdp_df = get_data_from_whoeurope(url_a)
                
                case "WHO/HESRI":
                    url_a = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
                    url_b = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
                    gdp_df = get_data_from_WHOHESR (url_code)
            
                case "WHO/HESRI 2":
                    url_a = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
                    url_b = "http://worldhealthorg.shinyapps.io/european_health_equity_dataset/"
                    gdp_df = get_data_from_WHOHESR (url_code)
            
            #filtering data by country selection
            dtc = gdp_df[(gdp_df['Country Code'].isin(iso_acronyms))]
            
            # Calculate the min and max of the 'Value' column
            value_min = dtc['Value'].min()
            value_max = dtc['Value'].max()
            
            # Calculate the center
            center = (value_min + value_max) / 2

            # Calculate the range for the y-axis
            y_min = center - (value_max - value_min)/2 * 1.2
            y_max = center + (value_max - value_min)/2 * 1.2
                      
            #preparing the chart command
            if (len(dtc) > 0):
                dtc = dtc.dropna()         
                column_titles = dtc.columns.tolist()
                dims_lower = [dim.lower() for dim in column_titles]
                print (url_code, column_titles, dims_lower)
                dims2 = [dim for dim in column_titles if ((dim != 'Year') and (dim != 'Value'))]
            
                #dims2 needed to draw dimensions lines in the chart
                if (len(dims2) > 1): dims2.remove("Country Code")

                #chart command
                chart = alt.Chart(dtc).mark_line(point=True).encode(
                x=alt.X('Year:Q', axis=alt.Axis(format='d')),   
                y=alt.Y('Value:Q', scale=alt.Scale(domain=[y_min, y_max])),
                color = "".join(dims2),                 
                tooltip=['Year', 'Value', 'Country Code']
                ).properties().interactive()
                        
                #Display the chart in Streamlit for each unique occurrence of dimension SEX
                c1.altair_chart(chart, use_container_width=True)
                
                #saving plot to Word file
                img_buffer = BytesIO()
                chart.save(img_buffer, format='png')  # Requires 'kaleido'
                img_buffer.seek(0)
                img_buffer.name = 'chart.png'  # python-docx expects .name attribute
                doc.add_picture(img_buffer, width=Inches(6))
            
                #pivoting data to show table below the chart
                pivot_data = pd.pivot_table(dtc, index= dims2, columns="Year", values="Value", aggfunc='sum').round(2)   
                c1.dataframe (pivot_data)
                doc.add_paragraph(pivot_data.to_string())
                
                #populating var prompt0 to feed AI 
                prompt0 += "\n" + chart_title + "\n" + pivot_data.to_string()
                
            else: 
                c1.write(":red[Data not available for this indicator]")
                
            c1.html("<a href=" + url_a + " target='_blank'>Data link...</a>")  
            c1.write ('*************************************************************')            
        
        print ("AI data")
        print ("Data for country:" + selected_country + prompt0)
        
        messages = []
        with st.spinner("Data analysis in progress... Country: "+ selected_country):
            
            #first prompt related to data analysis
            input ="You are an experienced data scientist. Use the following data for country " + selected_country + " and provide comment on trends and relations between indicators collected. Elaborate a concise report highlighting differences for males and females, for age groups, for groups with different education or income, for groups living in urban areas compared to rural, and trends over time" + prompt0 + " Do not add introductions or conclusions. No AI disclaimers or pleasantries. Use bullet points, titles and text."
            
            messages.append({"role": "user", "content": input})
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=messages
                )
            
            print ("Prompt data analysis" + input)
            assistant_reply = response.choices[0].message.content
            c1.write ("****** AI data report:")
            c1.write (assistant_reply)
            c1.write ("****************************************************************")
            doc.add_paragraph("****** AI data report:" + assistant_reply)
            
            # Second prompt for elaboration
            input2 = "You are a senior political advisor and you have to prepare a report with actionable points related to public health goods and services plan to be provided in your country. The country to focus is "    + selected_country + "." + selected_country + """ has a several laws and regulation related to health sector. Collect them for your reference. Your scope is to create a document with a different vision of health. Health is both a foundation and a goal of well-being economies. Health systems are not only economic sectors in their own right—employing millions and generating social value—but also key enablers of human development, social cohesion, and environmental sustainability. The vision will strengthen national capacities to generate, govern, and use health-related data to inform policies that promote equitable, resilient, and prosperous societies. 
            The vision works on 4 well-being capitals: Human well-being, Social well-being, Planetary well-being, Economic well-being. Human well-being is important because people's health and their subjective well-being are closely linked; both are drivers of economic prosperity, social mobility and cohesion. Human well-being is measured by indicators like: Healthy life expectancy, Mental health and well‑being, Ability to carry out daily activities free from illness, Universal Health Coverage, Quality and non‑discriminatory health & social care, Universal policies for housing food and fuel security, Early childhood development, Lifelong learning and literacy, Safe, orderly & regular migration. 
            Social well-being is represented by Trust, participation and social cohesion make significant contributions to mental and physical health and well-being, and are vital to building fair, peaceful, and resilient societies. Social well-being is measured by Living in safety and free from violence, Sense of belonging (“Mattering”), Social cohesion and embracing diversity, Perceived ability to influence politics and decisions, Social support and protection, Building trust in others and in institutions, Public spending on communities, Participation in volunteering.
            Planetary well-being is a key determinant of physical, mental and social well-being for current and future generations. It is also essential for economic prosperity. Environmental damage has significant negative impacts on well-being and prosperity. Planetary well-being is measured by Good air and water quality, Healthy and sustainable living environment, Sustainable public transport and active travel, Access to safe green space, Stable climate Biodiversity and natural capital, Circular economy and green technology.
            Economic well-being impacts physical and mental health and well-being and is essential to ensure that people have a sustainable income, as well as assets, so that they can prosper and participate in society. It's measured by Living wage, Universal social protection through the life‑course, Decent and psychologically safe work, Gender‑responsive employment, Social dialogue and collective bargaining, Economic cohesion and balanced development.
            
            Potential actions and activities carried by stakeholders to promote well-being are also mentioned in the WHO publications 'Health in the well-being economy' WHO/EURO:2023-7144-46910-68439 and in 'Deep dives on the well-being economy showcasing the experiences of Finland, Iceland, Scotland and Wales: summary of key findings' WHO/EURO:2023-7033-46799-68216. Please take inspiration by these publication whilst you prepare your report.
            Elaborate on the data given previously for """ + selected_country + """ around the 4 well-being capitals and indicators (where possible disaggregated by sex, gender, and age), if it's needed collect additional data and describe what are the key points to be considered for planning goods and services to promote the 4 well-being capitals.
            Provide a 5 page long report with actions for each well-being capital with content reference and data sources. Reason your points in relation to health laws and policies and data provided by highlighting data and relations among indicators. 
            Highlight data improvements that could be depending by the implementation of a law, expand key actions for each well-being capital by reasoning the choice with your comments and recommendations. At the end of your report make a summary table of recommendations, with well-being capitals as rows and key actions, legal/policy basis, expected impact as columns. 
            Do not add introductions or conclusions. No AI disclaimers or pleasantries. Use bullet points, titles and text."""
            
            messages.append({"role": "assistant", "content": assistant_reply})
            messages.append({"role": "user", "content": input2})

            print (messages)           
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=messages
                )
            assistant_reply = response.choices[0].message.content
            c1.write ("****** AI report:\n" + assistant_reply)
            doc.add_paragraph("****** AI report:" + assistant_reply)

            # --- Save doc to BytesIO and present download ---
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            c1.download_button(
                label="Download report as Word file",
                data=doc_buffer,
                file_name=selected_country+".docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            # -----------------------------------------------            
                )
  
else: 
    st.write ("Select to proceed...")              

#*********************************************************
            
    